"""
Local directory scanner — replaces zotero_client.py.
Scans a user-chosen references directory for PDFs plus RIS/BibTeX exports.
Subfolders become "collections".
"""

import hashlib
import json
import logging
import os
import re
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz

from app.crossref import fetch_crossref_metadata, merge_crossref_metadata, normalize_doi

logger = logging.getLogger(__name__)

PDF_EXTENSIONS = frozenset({".pdf"})
BIBLIOGRAPHY_EXTENSIONS = frozenset({".ris", ".bib", ".bibtex"})
TEXT_EXTENSIONS = frozenset({".txt"})
MARKDOWN_EXTENSIONS = frozenset({".md", ".markdown"})
CSV_EXTENSIONS = frozenset({".csv"})
WORD_EXTENSIONS = frozenset({".docx"})
IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif"})
SUPPORTED_EXTENSIONS = PDF_EXTENSIONS | BIBLIOGRAPHY_EXTENSIONS | TEXT_EXTENSIONS | MARKDOWN_EXTENSIONS | CSV_EXTENSIONS | WORD_EXTENSIONS | IMAGE_EXTENSIONS


def generate_item_key(file_path: str) -> str:
    return hashlib.md5(file_path.encode()).hexdigest()[:12]


def _generate_reference_item_key(file_path: str, entry_id: str, index: int) -> str:
    source = f"{file_path}::{entry_id or index}"
    return hashlib.md5(source.encode()).hexdigest()[:12]


def test_directory(dir_path: str) -> Dict[str, Any]:
    if not dir_path:
        return {"success": False, "item_count": 0, "message": "No directory path provided."}

    path = Path(dir_path).expanduser().resolve()
    if not path.exists():
        return {"success": False, "item_count": 0, "message": f"Directory does not exist: {path}"}
    if not path.is_dir():
        return {"success": False, "item_count": 0, "message": f"Not a directory: {path}"}

    files = list_files(path)
    folders = [d for d in path.iterdir() if d.is_dir() and not d.name.startswith(".")]

    msg = f"Connected. Found {len(files)} supported file(s) in {len(folders)} subfolder(s)."
    return {"success": True, "item_count": len(files), "folder_count": len(folders), "message": msg}


def list_files(root: Path) -> List[Path]:
    files = []
    for f in root.rglob("*"):
        if f.suffix.lower() in SUPPORTED_EXTENSIONS and not f.name.startswith("."):
            files.append(f)
    return files


def get_subfolders(root: Path) -> List[Dict[str, Any]]:
    """Return immediate subfolders only — used by the directory browser API."""
    folders = []
    try:
        for d in sorted(root.iterdir()):
            if d.is_dir() and not d.name.startswith("."):
                folders.append({
                    "collection_key": hashlib.md5(str(d).encode()).hexdigest()[:12],
                    "name": d.name,
                    "parent_key": "",
                    "path": str(d),
                    "rel_path": d.name,
                    "depth": 0,
                })
    except Exception as exc:
        logger.error("Error listing subfolders: %s", exc)
    return folders


MAX_SUBFOLDER_DEPTH = 10


def get_subfolders_recursive(root: Path) -> List[Dict[str, Any]]:
    """Recursively discover all subfolders with proper parent-child hierarchy."""
    folders: List[Dict[str, Any]] = []
    seen: set = set()

    def _walk(dir_path: Path, parent_key: str, rel_prefix: str, depth: int) -> None:
        if depth > MAX_SUBFOLDER_DEPTH:
            return
        real = str(dir_path.resolve())
        if real in seen:
            return
        seen.add(real)
        try:
            for d in sorted(dir_path.iterdir()):
                if d.is_dir() and not d.name.startswith("."):
                    col_key = hashlib.md5(str(d).encode()).hexdigest()[:12]
                    rel = f"{rel_prefix}/{d.name}" if rel_prefix else d.name
                    folders.append({
                        "collection_key": col_key,
                        "name": d.name,
                        "parent_key": parent_key,
                        "path": str(d),
                        "rel_path": rel,
                        "depth": depth,
                    })
                    _walk(d, col_key, rel, depth + 1)
        except Exception as exc:
            logger.error("Error listing subfolders of %s: %s", dir_path, exc)

    _walk(root, "", "", 0)
    return folders


def _clean_title(raw: str) -> str:
    title = raw.strip()
    title = re.sub(r"[\r\n]+", " ", title)
    if len(title) > 300:
        title = title[:300]
    return title


def _filename_to_title(filepath: Path) -> str:
    name = filepath.stem
    name = name.replace("_", " ").replace("-", " ")
    name = re.sub(r"\s+", " ", name).strip()
    if len(name) > 200:
        name = name[:200]
    return name


def _extract_year(text: str, metadata_date: str = "") -> str:
    if metadata_date:
        m = re.search(r"(\d{4})", metadata_date)
        if m:
            return m.group(1)
    if text:
        m = re.search(r"\b(19[5-9]\d|20[0-9]\d)\b", text[:2000])
        if m:
            return m.group(1)
    return ""


def _extract_doi(text: str) -> str:
    if not text:
        return ""
    m = re.search(r"(?:doi[:\s]*|https?://doi\.org/)(10\.\d{4,}/[^\s\"'>]+)", text, re.IGNORECASE)
    if m:
        return normalize_doi(m.group(1))
    return ""


def _extract_abstract(first_page_text: str) -> str:
    if not first_page_text:
        return ""

    patterns = [
        r"(?i)abstract[:\s]*\n?(.*?)(?:\n\s*(?:keywords|introduction|1[\.\s]|jel|resumen))",
        r"(?i)abstract[:\s]*(.*?)(?:\n\n)",
    ]

    for pattern in patterns:
        m = re.search(pattern, first_page_text, re.DOTALL)
        if m:
            abstract = m.group(1).strip()
            if len(abstract) > 50:
                if len(abstract) > 2000:
                    abstract = abstract[:2000]
                return abstract

    lines = first_page_text.strip().split("\n")
    for i, line in enumerate(lines[:20]):
        if re.match(r"(?i)^\s*abstract\b", line):
            abstract_parts = [re.sub(r"(?i)^\s*abstract[:\s]*", "", line)]
            for j in range(i + 1, min(i + 25, len(lines))):
                if re.match(r"(?i)^\s*(keywords|introduction|1[\.\s])", lines[j]):
                    break
                abstract_parts.append(lines[j])
            abstract = " ".join(abstract_parts).strip()
            if len(abstract) > 50:
                return abstract[:2000]

    return ""


def parse_authors(author_str: str) -> List[Dict]:
    if not author_str or not author_str.strip():
        return []

    creators = []
    separators = [";", " and ", " & ", ","]
    parts = [author_str]

    for sep in separators:
        new_parts = []
        for part in parts:
            new_parts.extend(part.split(sep))
        parts = new_parts

    for part in parts:
        name = part.strip()
        if not name:
            continue
        name = re.sub(r"\d+", "", name).strip()
        name = re.sub(r"\s+", " ", name).strip()
        if not name or len(name) < 2:
            continue

        if "," in name:
            last_first = name.split(",", 1)
            last_name = last_first[0].strip()
            first_name = last_first[1].strip() if len(last_first) > 1 else ""
            if last_name:
                creators.append({
                    "creatorType": "author",
                    "lastName": last_name,
                    "firstName": first_name,
                    "name": "",
                })
        else:
            name_parts = name.split()
            if len(name_parts) >= 2:
                creators.append({
                    "creatorType": "author",
                    "lastName": name_parts[-1],
                    "firstName": " ".join(name_parts[:-1]),
                    "name": "",
                })
            elif name_parts:
                creators.append({
                    "creatorType": "author",
                    "lastName": "",
                    "firstName": "",
                    "name": name,
                })

    return creators


def _parse_person_name(raw: str, creator_type: str = "author") -> Optional[Dict[str, str]]:
    name = re.sub(r"[{}]", "", raw or "").strip()
    name = re.sub(r"\s+", " ", name)
    if not name:
        return None

    if "," in name:
        last, first = [p.strip() for p in name.split(",", 1)]
        if last:
            return {
                "creatorType": creator_type,
                "lastName": last,
                "firstName": first,
                "name": "",
            }

    parts = name.split()
    if len(parts) >= 2:
        return {
            "creatorType": creator_type,
            "lastName": parts[-1],
            "firstName": " ".join(parts[:-1]),
            "name": "",
        }
    return {"creatorType": creator_type, "lastName": "", "firstName": "", "name": name}


def _parse_people(values: List[str], creator_type: str = "author") -> List[Dict[str, str]]:
    creators = []
    for value in values:
        for part in re.split(r"\s+\band\b\s+|;", value or "", flags=re.IGNORECASE):
            creator = _parse_person_name(part, creator_type)
            if creator:
                creators.append(creator)
    return creators


def _empty_reference_item(file_path: Path, item_key: str) -> Dict[str, Any]:
    try:
        from datetime import datetime
        date_modified = datetime.fromtimestamp(os.path.getmtime(str(file_path))).isoformat()
    except OSError:
        date_modified = ""

    return {
        "item_key": item_key,
        "file_path": str(file_path),
        "title": _filename_to_title(file_path),
        "creators": "[]",
        "year": "",
        "item_type": "journalArticle",
        "publication_title": "",
        "doi": "",
        "url": "",
        "abstract": "",
        "tags": "[]",
        "collection_keys": "[]",
        "date_modified": date_modified,
        "extra": "",
        "volume": "",
        "issue": "",
        "pages": "",
        "publisher": "",
        "place": "",
        "edition": "",
        "isbn": "",
        "issn": "",
        "citation_count": 0,
        "citation_count_updated_at": "",
        "full_text": "",
    }


def _map_ris_type(ris_type: str) -> str:
    mapping = {
        "JOUR": "journalArticle",
        "JFULL": "journalArticle",
        "CPAPER": "conferencePaper",
        "CONF": "conferencePaper",
        "BOOK": "book",
        "CHAP": "bookSection",
        "THES": "thesis",
        "RPRT": "report",
        "ELEC": "webpage",
        "WEB": "webpage",
        "NEWS": "newspaperArticle",
        "MGZN": "magazineArticle",
    }
    return mapping.get((ris_type or "").upper(), "journalArticle")


def _extract_reference_year(*values: str) -> str:
    for value in values:
        if not value:
            continue
        m = re.search(r"\b(19[5-9]\d|20[0-4]\d|202[0-6])\b", value)
        if m:
            return m.group(1)
    return ""


def _reference_full_text(item: Dict[str, Any]) -> str:
    parts = [
        item.get("title", ""),
        item.get("abstract", ""),
        item.get("publication_title", ""),
        item.get("doi", ""),
        item.get("url", ""),
        item.get("extra", ""),
    ]
    return "\n".join(p for p in parts if p)


def _read_text_file(file_path: Path) -> str:
    raw = file_path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def parse_ris_file(file_path: Path) -> List[Dict[str, Any]]:
    text = _read_text_file(file_path)
    records: List[Dict[str, List[str]]] = []
    current: Optional[Dict[str, List[str]]] = None
    current_tag = ""

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        tag_match = re.match(r"^([A-Z0-9]{2})\s{0,2}-\s?(.*)$", line)
        if tag_match:
            tag, value = tag_match.group(1).upper(), tag_match.group(2).strip()
            if tag == "TY":
                current = {"TY": [value]}
                records.append(current)
            elif current is not None:
                current.setdefault(tag, []).append(value)
            current_tag = tag
            continue

        if current is not None and current_tag and line.startswith(" ") and current.get(current_tag):
            current[current_tag][-1] = f"{current[current_tag][-1]} {line.strip()}".strip()

    items: List[Dict[str, Any]] = []
    for idx, record in enumerate(records, start=1):
        entry_id = (record.get("ID") or record.get("AN") or [str(idx)])[0]
        item = _empty_reference_item(file_path, _generate_reference_item_key(str(file_path), entry_id, idx))
        title = (record.get("TI") or record.get("T1") or record.get("CT") or record.get("BT") or [""])[0]
        publication = (record.get("JO") or record.get("JF") or record.get("JA") or record.get("T2") or [""])[0]
        start_page = (record.get("SP") or [""])[0]
        end_page = (record.get("EP") or [""])[0]
        pages = (record.get("M2") or [""])[0]
        if start_page and end_page:
            pages = f"{start_page}-{end_page}"
        elif start_page:
            pages = start_page

        creators = _parse_people(record.get("AU", []) + record.get("A1", []), "author")
        creators.extend(_parse_people(record.get("ED", []) + record.get("A2", []), "editor"))
        tags = [kw for kw in record.get("KW", []) if kw]

        import json
        item.update({
            "title": title or item["title"],
            "creators": json.dumps(creators),
            "year": _extract_reference_year(*(record.get("PY", []) + record.get("Y1", []) + record.get("DA", []))),
            "item_type": _map_ris_type((record.get("TY") or [""])[0]),
            "publication_title": publication,
            "doi": normalize_doi((record.get("DO") or [""])[0]),
            "url": (record.get("UR") or [""])[0],
            "abstract": (record.get("AB") or record.get("N2") or [""])[0],
            "tags": json.dumps(tags),
            "volume": (record.get("VL") or [""])[0],
            "issue": (record.get("IS") or [""])[0],
            "pages": pages,
            "publisher": (record.get("PB") or [""])[0],
            "place": (record.get("CY") or [""])[0],
        })

        sn = (record.get("SN") or [""])[0]
        if sn:
            if item["item_type"] == "book":
                item["isbn"] = sn
            else:
                item["issn"] = sn
        item["full_text"] = _reference_full_text(item)
        items.append(item)

    return items


def _map_bibtex_type(entry_type: str) -> str:
    mapping = {
        "article": "journalArticle",
        "book": "book",
        "inbook": "bookSection",
        "incollection": "bookSection",
        "inproceedings": "conferencePaper",
        "conference": "conferencePaper",
        "phdthesis": "thesis",
        "mastersthesis": "thesis",
        "techreport": "report",
        "report": "report",
        "online": "webpage",
        "misc": "webpage",
    }
    return mapping.get((entry_type or "").lower(), "journalArticle")


def _strip_bibtex_value(value: str) -> str:
    value = (value or "").strip()
    if (value.startswith("{") and value.endswith("}")) or (value.startswith('"') and value.endswith('"')):
        value = value[1:-1]
    value = re.sub(r"[{}]", "", value)
    value = value.replace("\\&", "&")
    return re.sub(r"\s+", " ", value).strip()


def _split_bibtex_fields(body: str) -> Dict[str, str]:
    fields: Dict[str, str] = {}
    i = 0
    length = len(body)
    while i < length:
        while i < length and body[i] in " \t\r\n,":
            i += 1
        key_start = i
        while i < length and re.match(r"[A-Za-z0-9_-]", body[i]):
            i += 1
        key = body[key_start:i].lower()
        while i < length and body[i].isspace():
            i += 1
        if not key or i >= length or body[i] != "=":
            i += 1
            continue
        i += 1
        while i < length and body[i].isspace():
            i += 1

        value_start = i
        if i < length and body[i] in "{":
            depth = 0
            while i < length:
                if body[i] == "{":
                    depth += 1
                elif body[i] == "}":
                    depth -= 1
                    if depth == 0:
                        i += 1
                        break
                i += 1
            value = body[value_start:i]
        elif i < length and body[i] == '"':
            i += 1
            escaped = False
            while i < length:
                if body[i] == '"' and not escaped:
                    i += 1
                    break
                escaped = body[i] == "\\" and not escaped
                if body[i] != "\\":
                    escaped = False
                i += 1
            value = body[value_start:i]
        else:
            while i < length and body[i] != ",":
                i += 1
            value = body[value_start:i]
        fields[key] = _strip_bibtex_value(value)
    return fields


def parse_bibtex_file(file_path: Path) -> List[Dict[str, Any]]:
    text = _read_text_file(file_path)
    entries: List[Tuple[str, str, str]] = []
    i = 0
    while i < len(text):
        at = text.find("@", i)
        if at < 0:
            break
        m = re.match(r"@([A-Za-z]+)\s*[{(]", text[at:])
        if not m:
            i = at + 1
            continue
        entry_type = m.group(1)
        open_pos = at + m.end() - 1
        close_char = "}" if text[open_pos] == "{" else ")"
        depth = 0
        j = open_pos
        while j < len(text):
            if text[j] == text[open_pos]:
                depth += 1
            elif text[j] == close_char:
                depth -= 1
                if depth == 0:
                    break
            j += 1
        content = text[open_pos + 1:j]
        comma = content.find(",")
        if comma >= 0:
            entries.append((entry_type, content[:comma].strip(), content[comma + 1:]))
        i = j + 1

    items: List[Dict[str, Any]] = []
    for idx, (entry_type, cite_key, body) in enumerate(entries, start=1):
        fields = _split_bibtex_fields(body)
        item = _empty_reference_item(file_path, _generate_reference_item_key(str(file_path), cite_key, idx))
        creators = _parse_people([fields.get("author", "")], "author")
        creators.extend(_parse_people([fields.get("editor", "")], "editor"))
        keywords = [k.strip() for k in re.split(r"[,;]", fields.get("keywords", "")) if k.strip()]

        import json
        item.update({
            "title": fields.get("title") or item["title"],
            "creators": json.dumps(creators),
            "year": _extract_reference_year(fields.get("year", ""), fields.get("date", "")),
            "item_type": _map_bibtex_type(entry_type),
            "publication_title": fields.get("journal") or fields.get("journaltitle") or fields.get("booktitle") or "",
            "doi": normalize_doi(fields.get("doi", "")),
            "url": fields.get("url", ""),
            "abstract": fields.get("abstract", ""),
            "tags": json.dumps(keywords),
            "volume": fields.get("volume", ""),
            "issue": fields.get("number") or fields.get("issue", ""),
            "pages": fields.get("pages", "").replace("--", "-"),
            "publisher": fields.get("publisher", ""),
            "place": fields.get("address") or fields.get("location", ""),
            "edition": fields.get("edition", ""),
            "isbn": fields.get("isbn", ""),
            "issn": fields.get("issn", ""),
            "extra": f"BibTeX key: {cite_key}" if cite_key else "",
        })
        item["full_text"] = _reference_full_text(item)
        items.append(item)

    return items


def extract_reference_metadata(file_path: Path) -> List[Dict[str, Any]]:
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".ris":
            return parse_ris_file(file_path)
        if suffix in {".bib", ".bibtex"}:
            return parse_bibtex_file(file_path)
    except Exception as exc:
        logger.error("Error parsing bibliography file %s: %s", file_path, exc)
    return []


_REFERENCE_OVERLAY_FIELDS = (
    "title",
    "creators",
    "year",
    "item_type",
    "publication_title",
    "doi",
    "url",
    "abstract",
    "volume",
    "issue",
    "pages",
    "publisher",
    "place",
    "edition",
    "isbn",
    "issn",
)


def _normalize_match_text(value: str) -> str:
    value = (value or "").lower()
    value = re.sub(r"\bdoi\s*:?\s*10\.\S+", " ", value)
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _similarity(a: str, b: str) -> float:
    a_norm = _normalize_match_text(a)
    b_norm = _normalize_match_text(b)
    if not a_norm or not b_norm:
        return 0.0
    return SequenceMatcher(None, a_norm, b_norm).ratio()


def _json_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if not value:
        return []
    try:
        parsed = json.loads(value) if isinstance(value, str) else value
        return parsed if isinstance(parsed, list) else []
    except (json.JSONDecodeError, TypeError):
        return []


def _creator_last_names(item: Dict[str, Any]) -> List[str]:
    names = []
    for creator in _json_list(item.get("creators", "[]")):
        if not isinstance(creator, dict):
            continue
        last = creator.get("lastName") or creator.get("last_name") or creator.get("name") or ""
        last = _normalize_match_text(str(last))
        if last:
            names.append(last)
    return names


def _reference_pdf_match_score(reference_item: Dict[str, Any], file_item: Dict[str, Any]) -> float:
    ref_doi = normalize_doi(reference_item.get("doi", ""))
    pdf_doi = normalize_doi(file_item.get("doi", ""))
    if ref_doi and pdf_doi and ref_doi == pdf_doi:
        return 100.0

    ref_title = reference_item.get("title", "")
    pdf_title = file_item.get("title", "")
    pdf_stem = Path(file_item.get("file_path", "")).stem
    best_title_score = max(
        _similarity(ref_title, pdf_title),
        _similarity(ref_title, pdf_stem),
    ) * 80

    ref_year = str(reference_item.get("year", "") or "")
    pdf_year = str(file_item.get("year", "") or "")
    year_bonus = 8 if ref_year and (ref_year == pdf_year or ref_year in pdf_stem) else 0

    filename_norm = _normalize_match_text(pdf_stem)
    creator_bonus = 0
    for last_name in _creator_last_names(reference_item)[:3]:
        if last_name and last_name in filename_norm:
            creator_bonus = 8
            break

    return best_title_score + year_bonus + creator_bonus


def _merge_reference_into_file_item(
    file_item: Dict[str, Any],
    reference_item: Dict[str, Any],
) -> Dict[str, Any]:
    merged = dict(file_item)
    for field in _REFERENCE_OVERLAY_FIELDS:
        value = reference_item.get(field)
        if value not in (None, "", "[]"):
            merged[field] = value

    pdf_tags = _json_list(file_item.get("tags", "[]"))
    ref_tags = _json_list(reference_item.get("tags", "[]"))
    if pdf_tags or ref_tags:
        merged_tags = []
        seen = set()
        for tag in pdf_tags + ref_tags:
            key = str(tag).strip().lower()
            if key and key not in seen:
                merged_tags.append(tag)
                seen.add(key)
        merged["tags"] = json.dumps(merged_tags)

    # Keep the real document identity and extracted content from the file item.
    merged["item_key"] = file_item["item_key"]
    merged["file_path"] = file_item["file_path"]
    merged["date_modified"] = file_item.get("date_modified", "")
    merged["full_text"] = file_item.get("full_text", "")
    merged["_merged_reference_item_keys"] = [reference_item["item_key"]]
    return merged


def _extract_authors_from_first_page(text: str) -> List[Dict]:
    """
    Extract author names from the first page of an academic PDF.
    """
    if not text or len(text) < 50:
        return []

    lines = [l.strip() for l in text.split("\n") if l.strip()]

    # Find the title line (first long line that looks like an article title)
    title_idx = -1
    for i, line in enumerate(lines[:15]):
        if len(line) < 20:
            continue
        if re.match(r"^https?://", line):
            continue
        if re.match(r"^file://", line, re.IGNORECASE):
            continue
        if re.match(r"^\d{2}/\d{2}/\d{2}", line):  # browser print date stamp
            continue
        if re.match(r"(?i)^page\s+\d+\s+of\s+\d+", line):
            continue
        if re.match(r"^\d{4}-\d+", line):
            continue
        if re.match(r"(?i)^(available\s+online|received|accepted|published|copyright|©|http)", line):
            continue
        # Skip journal header lines like "Science of the Total Environment 947 (2024) 174530"
        if re.search(r"\d{4}\)\s*\d{4,7}$", line):
            continue
        # Skip lines that look like journal headers: "Journal Name Vol (Year) Pages"
        if re.match(r"^[A-Z][a-z].*?\d+\s*\(\d{4}\)", line):
            continue
        title_idx = i
        break

    if title_idx < 0:
        return []

    # Collect author lines right after the title
    author_line = ""
    for j in range(title_idx + 1, min(title_idx + 5, len(lines))):
        line = lines[j]
        if re.match(r"(?i)^(abstract|keywords|introduction|highlights|graphical|article|received|available|editor|history)", line):
            break
        if re.match(r"^file://", line, re.IGNORECASE):
            break
        if re.match(r"(?i)^page\s+\d+\s+of\s+\d+", line):
            break
        if re.match(r"^\d{2}/\d{2}/\d{2}", line):
            break
        if re.search(r"(?i)(university|institute|department|school\s+of|college|faculty|laboratory|research\s+(?:center|institute|station)|corresponding\s+author)", line) and len(line) > 25:
            break
        if line.startswith(("a ", "b ", "c ", "d ")) and re.search(r"(?i)(university|institute|department|school)", line):
            break
        author_line += " " + line

    author_line = author_line.strip()
    if not author_line or len(author_line) < 5:
        return []

    # Clean: "Kesheng Huang a, Jinfeng Wu a, Zhengxiao Fu a, Jianhui Du a,b,*"
    # Remove superscript markers: " a", " a,b", " a,b,*"
    cleaned = re.sub(r"\s+[a-z](?:,[a-z])*(?:,\*)?(?=\s*,|\s*$)", "", author_line)
    cleaned = re.sub(r"\s*\*", "", cleaned)
    cleaned = re.sub(r"\s+,", ",", cleaned)
    cleaned = re.sub(r",\s+", ", ", cleaned)
    cleaned = cleaned.strip().rstrip(",").strip()

    parts = [p.strip() for p in cleaned.split(",") if p.strip()]

    creators = []
    for part in parts:
        words = part.split()
        if len(words) >= 2 and all(w[0].isupper() for w in words):
            creators.append({
                "creatorType": "author",
                "lastName": words[-1],
                "firstName": " ".join(words[:-1]),
                "name": "",
            })

    return creators if len(creators) >= 2 else []


def extract_pdf_fulltext(pdf_path: Path) -> str:
    """Extract full text from a PDF without re-parsing metadata. Called lazily during sync."""
    try:
        doc = fitz.open(str(pdf_path))
        try:
            return "".join(page.get_text() + "\n" for page in doc)
        finally:
            doc.close()
    except Exception as exc:
        logger.warning("Could not extract text from %s: %s", pdf_path, exc)
        return ""


def extract_pdf_metadata(pdf_path: Path, extract_text: bool = True) -> Dict[str, Any]:
    item_key = generate_item_key(str(pdf_path))
    result: Dict[str, Any] = {
        "item_key": item_key,
        "file_path": str(pdf_path),
        "title": _filename_to_title(pdf_path),
        "creators": "[]",
        "year": "",
        "item_type": "journalArticle",
        "publication_title": "",
        "doi": "",
        "url": "",
        "abstract": "",
        "tags": "[]",
        "collection_keys": "[]",
        "date_modified": "",
        "extra": "",
        "volume": "",
        "issue": "",
        "pages": "",
        "publisher": "",
        "place": "",
        "edition": "",
        "isbn": "",
        "issn": "",
        "full_text": "",
    }

    try:
        doc = fitz.open(str(pdf_path))
    except Exception as exc:
        logger.error("Cannot open PDF %s: %s", pdf_path, exc)
        return result

    try:
        meta = doc.metadata or {}

        raw_title = meta.get("title", "") or ""
        if raw_title.strip():
            result["title"] = _clean_title(raw_title)

        raw_author = meta.get("author", "") or ""
        creators = parse_authors(raw_author)

        first_page_text = doc[0].get_text() if doc.page_count > 0 else ""

        if len(creators) <= 1:
            page_creators = _extract_authors_from_first_page(first_page_text or "")
            if len(page_creators) > len(creators):
                creators = page_creators

        import json
        result["creators"] = json.dumps(creators)

        full_text = ""
        if extract_text:
            try:
                for page in doc:
                    full_text += page.get_text() + "\n"
            except Exception as exc:
                logger.warning("Could not extract text from %s: %s", pdf_path, exc)

        result["full_text"] = full_text

        result["year"] = _extract_year(
            first_page_text[:2000],
            meta.get("creationDate", "") or "",
        )

        result["doi"] = _extract_doi(first_page_text[:5000])

        # Extract journal info from subject field (Elsevier puts it there)
        subject = meta.get("subject", "") or ""
        if subject:
            if not result["abstract"]:
                result["abstract"] = subject[:2000]
            # Elsevier subject format: "Journal Name, Vol (Year) Pages. doi:..."
            subj_journal = re.match(r"^([A-Z][^,]+?),\s*(\d+)\s*\((\d{4})\)\s*([\d]+)", subject)
            if subj_journal:
                if not result["publication_title"]:
                    result["publication_title"] = subj_journal.group(1).strip()
                if not result["volume"]:
                    result["volume"] = subj_journal.group(2)
                if not result["pages"]:
                    result["pages"] = subj_journal.group(4)

        abstract = _extract_abstract(first_page_text)
        if abstract:
            result["abstract"] = abstract

        result["item_type"] = _detect_pdf_item_type(meta, first_page_text or "", bool(result["doi"]))

        if extract_text and not result["abstract"] and result["item_type"] == "document":
            summary = _extract_document_summary(full_text[:3000])
            if summary:
                result["abstract"] = summary

        keywords = meta.get("keywords", "") or ""
        if keywords:
            tags = [k.strip() for k in keywords.split(",") if k.strip()]
            result["tags"] = json.dumps(tags)

        mod_date = meta.get("modDate", "") or ""
        if mod_date:
            result["date_modified"] = mod_date
        else:
            mtime = os.path.getmtime(str(pdf_path))
            from datetime import datetime
            result["date_modified"] = datetime.fromtimestamp(mtime).isoformat()

        _extract_journal_info(first_page_text or "", result)

        if result["doi"]:
            crossref_metadata = fetch_crossref_metadata(result["doi"])
            if crossref_metadata:
                result = merge_crossref_metadata(result, crossref_metadata)
                logger.info("Applied Crossref metadata for DOI %s", result["doi"])

    except Exception as exc:
        logger.error("Error extracting metadata from %s: %s", pdf_path, exc)
    finally:
        doc.close()

    return result


def _extract_journal_info(text: str, result: Dict[str, Any]) -> None:
    """
    Extract journal name, volume, issue, pages, article number, ISSN from
    the first page of an academic PDF. Covers major publishers:
    Elsevier, Springer, Wiley, Taylor & Francis, MDPI, etc.
    """
    if not text:
        return

    header = text[:4000]

    # ── Journal name ──
    journal_patterns = [
        r"(?i)Science\s+of\s+the\s+Total\s+Environment",
        r"(?i)Journal\s+of\s+[A-Z][a-zA-Z\s]+",
        r"(?i)(?:Published\s+in|journal\s*(?:of|for)?)\s*:?\s*([A-Z][A-Za-z\s&]+?)(?:\n|,|\d|Vol\.?)",
        r"(?:^|\n)\s*([A-Z][a-z]+(?:\s+(?:of|for|the|in|and|&|a|an)\s+[A-Z][a-zA-Z\s&]+?))\s*(?:\n|Vol|\d{4})",
    ]

    known_journals = {
        "science of the total environment": "Science of the Total Environment",
        "journal of hydrology": "Journal of Hydrology",
        "water resources research": "Water Resources Research",
        "agricultural and forest meteorology": "Agricultural and Forest Meteorology",
        "environmental research letters": "Environmental Research Letters",
        "journal of climate": "Journal of Climate",
        "climate dynamics": "Climate Dynamics",
        "international journal of climatology": "International Journal of Climatology",
        "theoretical and applied climatology": "Theoretical and Applied Climatology",
        "remote sensing of environment": "Remote Sensing of Environment",
        "journal of environmental management": "Journal of Environmental Management",
        "environmental modelling & software": "Environmental Modelling & Software",
        "catena": "Catena",
        "geoderma": "Geoderma",
        "soil science society of america journal": "Soil Science Society of America Journal",
        "nature climate change": "Nature Climate Change",
        "global change biology": "Global Change Biology",
        "ecological indicators": "Ecological Indicators",
        "stochastic environmental research and risk assessment": "Stochastic Environmental Research and Risk Assessment",
    }

    text_lower = text[:6000].lower()
    for key, name in known_journals.items():
        if key in text_lower:
            result["publication_title"] = name
            break

    if not result["publication_title"]:
        for pattern in journal_patterns:
            m = re.search(pattern, header)
            if m:
                journal = m.group(1) if m.lastindex else m.group(0)
                journal = journal.strip()
                if 5 < len(journal) < 120 and not re.match(r"(?i)(abstract|introduction|keywords)", journal):
                    result["publication_title"] = journal
                    break

    # ── Volume / Issue / Pages ──
    vol_patterns = [
        r"(?i)Volume\s+(\d+)",
        r"(?i)Vol\.?\s*(\d+)",
        r"(?i)V(?:ol)?\.?\s*(\d{1,5})\s*,?\s*(?:No|Issue|n)\.?\s*(\d+)",
    ]
    for pat in vol_patterns:
        m = re.search(pat, header)
        if m:
            result["volume"] = m.group(1)
            if m.lastindex and m.lastindex >= 2:
                result["issue"] = m.group(2)
            break

    # Article number (common in Elsevier: "174530")
    art_match = re.search(r"(?:Article\s+(?:number|ID)|Article\s+)?\b(\d{5,7})\b", header[:2000])
    if art_match and not result["pages"]:
        num = art_match.group(1)
        if num.startswith("1") and len(num) in (5, 6, 7):
            result["pages"] = num

    # Pages: pp. 123-145 or 123–145
    pages_match = re.search(r"(?:pp\.?\s*|pages?\s*)(\d{1,5})\s*[-–]\s*(\d{1,5})", header, re.IGNORECASE)
    if pages_match:
        result["pages"] = f"{pages_match.group(1)}-{pages_match.group(2)}"

    # ── ISSN ──
    issn_match = re.search(r"(?:ISSN|issn)[:\s]*(\d{4}-\d{3}[\dX])", header)
    if issn_match:
        result["issn"] = issn_match.group(1)

    # ── Publisher ──
    pub_match = re.search(r"(?i)(Elsevier|Springer|Wiley|Taylor\s*&\s*Francis|MDPI|Oxford\s+University|Cambridge\s+University|IEEE)", header[:3000])
    if pub_match:
        result["publisher"] = pub_match.group(1)


# ── PDF annotation import ─────────────────────────────────────────────────────

# fitz annotation type code → app annotation_type string (None = skip)
_FITZ_TYPE_MAP: Dict[int, Optional[str]] = {
    0:  "comment",    # Text / sticky note
    2:  "comment",    # FreeText
    4:  "area",       # Square
    5:  "area",       # Circle
    6:  "area",       # Polygon
    8:  "highlight",  # Highlight
    9:  "underline",  # Underline
    10: "underline",  # Squiggly
    11: "underline",  # StrikeOut (closest visual match)
    19: "comment",    # Caret
    # 3=Line, 14=Ink, 12=Stamp → None (skipped)
}


def _fitz_color_to_hex(color_tuple) -> str:
    if not color_tuple or len(color_tuple) < 3:
        return "#ffff00"
    r, g, b = color_tuple[:3]
    return "#{:02x}{:02x}{:02x}".format(int(round(r * 255)), int(round(g * 255)), int(round(b * 255)))


def _quads_to_norm_rects(vertices, pw: float, ph: float) -> List[Dict]:
    """Convert fitz quad vertex list (4 Points per quad) to normalized {x,y,width,height} dicts."""
    rects = []
    for i in range(0, len(vertices), 4):
        group = vertices[i:i + 4]
        if len(group) < 4:
            break
        xs = [p.x for p in group]
        ys = [p.y for p in group]
        x0, y0, x1, y1 = min(xs), min(ys), max(xs), max(ys)
        if x1 - x0 < 0.5 or y1 - y0 < 0.5:
            continue
        rects.append({
            "x": x0 / pw,
            "y": y0 / ph,
            "width": (x1 - x0) / pw,
            "height": (y1 - y0) / ph,
        })
    return rects


def _rect_to_norm_rects(rect, pw: float, ph: float) -> List[Dict]:
    x0, y0, x1, y1 = rect.x0, rect.y0, rect.x1, rect.y1
    if x1 - x0 < 0.5 or y1 - y0 < 0.5:
        return []
    return [{"x": x0 / pw, "y": y0 / ph, "width": (x1 - x0) / pw, "height": (y1 - y0) / ph}]


def _quoted_text_from_page(page, vertices, rect) -> str:
    """Extract the text covered by highlight quads, falling back to the bounding rect."""
    parts = []
    try:
        if vertices:
            for i in range(0, len(vertices), 4):
                group = vertices[i:i + 4]
                if len(group) < 4:
                    break
                xs = [p.x for p in group]
                ys = [p.y for p in group]
                clip = fitz.Rect(min(xs), min(ys), max(xs), max(ys))
                text = page.get_textbox(clip).strip()
                if text:
                    parts.append(text)
        if not parts and rect:
            text = page.get_textbox(rect).strip()
            if text:
                parts.append(text)
    except Exception:
        pass
    return "\n".join(parts)


def extract_pdf_annotations(pdf_path: Path) -> List[Dict[str, Any]]:
    """
    Read all importable annotations embedded in a PDF.

    Returns a list of dicts matching the annotations table columns
    (without annotation_id / file_id / created_at / updated_at).
    source_chunk_id is set to 'imported:<hash>' for deduplication.
    """
    import hashlib as _hashlib
    import json as _json

    results: List[Dict[str, Any]] = []
    try:
        doc = fitz.open(str(pdf_path))
    except Exception as exc:
        logger.error("Cannot open PDF for annotation import %s: %s", pdf_path, exc)
        return results

    try:
        for page in doc:
            pw = page.rect.width
            ph = page.rect.height
            if pw == 0 or ph == 0:
                continue
            page_idx = page.number  # 0-based, matches app's page_index

            for annot in page.annots():
                type_code = annot.type[0]
                app_type = _FITZ_TYPE_MAP.get(type_code)
                if app_type is None:
                    continue

                # ── Color ────────────────────────────────────────────────────
                colors = annot.colors or {}
                raw_color = colors.get("stroke") or colors.get("fill")
                color = _fitz_color_to_hex(raw_color)

                # ── Geometry ─────────────────────────────────────────────────
                vertices = annot.vertices or []
                text_mark = type_code in (8, 9, 10, 11)  # highlight / underline family
                if text_mark and vertices:
                    rects = _quads_to_norm_rects(vertices, pw, ph)
                else:
                    rects = _rect_to_norm_rects(annot.rect, pw, ph)

                if not rects:
                    continue

                geometry_json = _json.dumps({"rects": rects})

                # ── Text content ─────────────────────────────────────────────
                info = annot.info or {}
                content = (info.get("content") or "").strip()

                if text_mark:
                    # Extract actual text from the page; some apps (Zotero) also
                    # put it in content — prefer page extraction, fall back to content
                    quote = _quoted_text_from_page(page, vertices if vertices else None, annot.rect)
                    if not quote:
                        quote = content
                        content = ""
                    comment = content  # whatever remains is a note on the highlight
                elif app_type == "comment":
                    quote = ""
                    comment = content
                else:
                    quote = ""
                    comment = content

                # ── Deduplication hash ───────────────────────────────────────
                src_hash = _hashlib.md5(
                    f"{page_idx}|{app_type}|{geometry_json}|{quote[:120]}".encode()
                ).hexdigest()[:16]

                results.append({
                    "page_index": page_idx,
                    "annotation_type": app_type,
                    "color": color,
                    "quote": quote,
                    "comment": comment,
                    "geometry_json": geometry_json,
                    "source_chunk_id": f"imported:{src_hash}",
                    "sentiment": None,
                })
    except Exception as exc:
        logger.error("Error reading annotations from %s: %s", pdf_path, exc)
    finally:
        doc.close()

    return results


_BROWSER_CREATORS = frozenset({"safari", "chrome", "firefox", "webkit", "mozilla"})
_BROWSER_PRODUCERS = frozenset({"quartz pdfcontext", "skia", "pdfium", "chromium"})
_WORD_PROCESSOR_CREATORS = frozenset({"microsoft word", "libreoffice", "pages", "google docs", "powerpoint", "keynote"})


def _detect_pdf_item_type(meta: dict, first_page_text: str, has_doi: bool) -> str:
    """Return the most appropriate item_type for a PDF based on its producer metadata."""
    if has_doi:
        return "journalArticle"

    creator = (meta.get("creator") or "").lower()
    producer = (meta.get("producer") or "").lower()

    is_browser = (
        any(b in creator for b in _BROWSER_CREATORS) or
        any(b in producer for b in _BROWSER_PRODUCERS) or
        bool(first_page_text and re.search(r"^file:///", first_page_text, re.MULTILINE))
    )
    if is_browser:
        return "document"

    is_word_processor = (
        any(w in creator for w in _WORD_PROCESSOR_CREATORS) or
        any(w in producer for w in _WORD_PROCESSOR_CREATORS)
    )
    if is_word_processor:
        return "document"

    return "journalArticle"


def _extract_document_summary(text: str, max_chars: int = 400) -> str:
    """First meaningful paragraph from a non-academic document (browser prints, notes, etc.)."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    collected = []
    for line in lines:
        if re.match(r"^(file://|https?://)", line, re.IGNORECASE):
            continue
        if re.match(r"(?i)^page\s+\d+\s+of\s+\d+", line):
            continue
        if re.match(r"^\d{2}/\d{2}/\d{2}", line):
            continue
        if len(line) < 30:
            continue
        collected.append(line)
        if sum(len(l) for l in collected) >= max_chars:
            break
    return " ".join(collected)[:max_chars].strip()


def extract_image_metadata(file_path: Path) -> Dict[str, Any]:
    item = _base_item(file_path, "image")
    try:
        from PIL import Image, ExifTags
        img = Image.open(str(file_path))
        w, h = img.size
        fmt = img.format or file_path.suffix.upper().lstrip(".")
        item["abstract"] = f"{fmt} · {w}×{h} px"
        try:
            exif = img._getexif() if hasattr(img, "_getexif") else None
            if exif:
                for tag_id, val in exif.items():
                    tag_name = ExifTags.TAGS.get(tag_id, "")
                    if tag_name in ("DateTimeOriginal", "DateTime") and val:
                        m = re.match(r"(\d{4})", str(val))
                        if m:
                            item["year"] = m.group(1)
                            break
        except Exception:
            pass
        img.close()
    except ImportError:
        logger.debug("Pillow not available — image metadata will be filename-only")
    except Exception as exc:
        logger.warning("Cannot read image %s: %s", file_path, exc)
    return item


def _base_item(file_path: Path, item_type: str = "document") -> Dict[str, Any]:
    """Shared base dict for non-PDF items."""
    import os
    mtime = os.path.getmtime(str(file_path))
    import datetime
    year = str(datetime.datetime.fromtimestamp(mtime).year)
    return {
        "item_key": generate_item_key(str(file_path)),
        "file_path": str(file_path),
        "title": _filename_to_title(file_path),
        "creators": "[]",
        "year": year,
        "item_type": item_type,
        "publication_title": "",
        "doi": "",
        "url": "",
        "abstract": "",
        "tags": "[]",
        "collection_keys": "[]",
        "date_modified": "",
        "extra": "",
        "volume": "",
        "issue": "",
        "pages": "",
        "publisher": "",
        "place": "",
        "edition": "",
        "isbn": "",
        "issn": "",
        "full_text": "",
    }


def extract_txt_metadata(file_path: Path) -> Dict[str, Any]:
    item = _base_item(file_path, "document")
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        item["full_text"] = text
        # Use first non-empty line as title if it looks like a title
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        if lines and len(lines[0]) < 120:
            item["title"] = lines[0]
        # Abstract = first 500 chars of remaining text
        rest = "\n".join(lines[1:]) if len(lines) > 1 else ""
        item["abstract"] = rest[:500]
    except Exception as exc:
        logger.warning("Cannot read txt %s: %s", file_path, exc)
    return item


def extract_md_metadata(file_path: Path) -> Dict[str, Any]:
    item = _base_item(file_path, "note")
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        item["full_text"] = text

        # Parse YAML frontmatter if present
        import re as _re
        fm_match = _re.match(r'^---\s*\n(.*?)\n---\s*\n', text, _re.DOTALL)
        body = text
        if fm_match:
            fm_text = fm_match.group(1)
            body = text[fm_match.end():]
            for line in fm_text.splitlines():
                if ':' in line:
                    k, v = line.split(':', 1)
                    k, v = k.strip().lower(), v.strip().strip('"\'')
                    if k == 'title' and v:
                        item["title"] = v
                    elif k in ('author', 'authors') and v:
                        item["creators"] = _parse_md_authors(v)
                    elif k in ('date', 'year') and v:
                        item["year"] = v[:4]
                    elif k in ('abstract', 'description') and v:
                        item["abstract"] = v

        # If no title from frontmatter, use first # heading
        if item["title"] == _filename_to_title(file_path):
            h_match = _re.search(r'^#{1,3}\s+(.+)$', body, _re.MULTILINE)
            if h_match:
                item["title"] = h_match.group(1).strip()

        # Abstract from first paragraph if not set
        if not item["abstract"]:
            paras = [p.strip() for p in _re.split(r'\n{2,}', body) if p.strip()]
            # Skip heading paragraphs
            for p in paras:
                if not p.startswith('#'):
                    item["abstract"] = p[:500]
                    break

    except Exception as exc:
        logger.warning("Cannot read md %s: %s", file_path, exc)
    return item


def _parse_md_authors(v: str) -> str:
    import json
    # Handle "FirstName LastName, FirstName LastName" or "[name1, name2]"
    v = v.strip('[]')
    parts = [p.strip() for p in v.split(',') if p.strip()]
    creators = []
    for p in parts:
        words = p.strip().split()
        if len(words) >= 2:
            creators.append({"firstName": " ".join(words[:-1]), "lastName": words[-1]})
        elif words:
            creators.append({"name": words[0]})
    return json.dumps(creators)


def extract_csv_metadata(file_path: Path) -> Dict[str, Any]:
    item = _base_item(file_path, "dataset")
    try:
        import csv as _csv
        text = file_path.read_text(encoding="utf-8", errors="replace")
        item["full_text"] = text
        reader = _csv.reader(text.splitlines())
        rows = list(reader)
        if rows:
            headers = rows[0]
            item["abstract"] = f"Dataset with {len(headers)} columns: {', '.join(headers[:10])}. {len(rows)-1} rows."
        else:
            item["abstract"] = "Empty CSV file."
    except Exception as exc:
        logger.warning("Cannot read csv %s: %s", file_path, exc)
    return item


def extract_docx_metadata(file_path: Path) -> Dict[str, Any]:
    item = _base_item(file_path, "document")
    try:
        from docx import Document
        import json
        doc = Document(str(file_path))

        # Core properties
        props = doc.core_properties
        if props.title:
            item["title"] = props.title.strip()
        if props.author:
            parts = props.author.strip().split()
            if len(parts) >= 2:
                item["creators"] = json.dumps([{"firstName": " ".join(parts[:-1]), "lastName": parts[-1]}])
            elif parts:
                item["creators"] = json.dumps([{"name": props.author.strip()}])
        if props.created:
            item["year"] = str(props.created.year)

        # Extract full text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        item["full_text"] = full_text

        # If no title from props, use first non-empty paragraph or heading
        if item["title"] == _filename_to_title(file_path):
            for p in doc.paragraphs:
                t = p.text.strip()
                if t and len(t) < 150:
                    item["title"] = t
                    break

        # Abstract: first ~500 chars of body (skip title paragraph)
        body_paras = [p for p in paragraphs[1:] if len(p) > 20]
        if body_paras:
            item["abstract"] = body_paras[0][:500]

    except Exception as exc:
        logger.warning("Cannot read docx %s: %s", file_path, exc)
    return item


def scan_directory(dir_path: str, progress_callback=None) -> Tuple[List[Dict], List[Dict]]:
    root = Path(dir_path).expanduser().resolve()
    if not root.exists() or not root.is_dir():
        raise ValueError(f"Invalid directory: {dir_path}")

    collections = get_subfolders_recursive(root)
    folder_map = {c["path"]: c["collection_key"] for c in collections}

    files = sorted(list_files(root), key=lambda p: str(p).lower())
    reference_entries_by_file: Dict[Path, List[Dict[str, Any]]] = {}
    reference_entries_by_folder: Dict[Path, List[Dict[str, Any]]] = {}
    for bibliography_file in files:
        if bibliography_file.suffix.lower() not in BIBLIOGRAPHY_EXTENSIONS:
            continue
        entries = extract_reference_metadata(bibliography_file)
        if entries:
            reference_entries_by_file[bibliography_file] = entries
            reference_entries_by_folder.setdefault(bibliography_file.parent, []).extend(entries)

    matched_reference_keys = set()
    total = len(files)
    items = []

    def collection_keys_for(file_path: Path) -> List[str]:
        parent = file_path.parent
        keys = []
        while parent != root and str(parent).startswith(str(root)):
            if str(parent) in folder_map:
                keys.append(folder_map[str(parent)])
            parent = parent.parent
        return keys or ["root"]

    for idx, file_path in enumerate(files):
        if progress_callback and idx % 5 == 0:
            progress_callback(f"Scanning {idx + 1}/{total}", file_path.name[:55])

        if file_path.suffix.lower() in PDF_EXTENSIONS:
            file_items = [extract_pdf_metadata(file_path, extract_text=False)]
            references = reference_entries_by_folder.get(file_path.parent, [])
            for file_item in file_items:
                best_reference = None
                best_score = 0.0
                for reference_item in references:
                    if reference_item["item_key"] in matched_reference_keys:
                        continue
                    score = _reference_pdf_match_score(reference_item, file_item)
                    if score > best_score:
                        best_reference = reference_item
                        best_score = score
                if best_reference and best_score >= 58:
                    file_item.update(_merge_reference_into_file_item(file_item, best_reference))
                    matched_reference_keys.add(best_reference["item_key"])
        elif file_path.suffix.lower() in IMAGE_EXTENSIONS:
            file_items = [extract_image_metadata(file_path)]
        elif file_path.suffix.lower() in TEXT_EXTENSIONS:
            file_items = [extract_txt_metadata(file_path)]
        elif file_path.suffix.lower() in MARKDOWN_EXTENSIONS:
            file_items = [extract_md_metadata(file_path)]
        elif file_path.suffix.lower() in CSV_EXTENSIONS:
            file_items = [extract_csv_metadata(file_path)]
        elif file_path.suffix.lower() in WORD_EXTENSIONS:
            file_items = [extract_docx_metadata(file_path)]
        else:
            continue

        col_keys = collection_keys_for(file_path)
        for item_data in file_items:
            item_data["collection_keys"] = json.dumps(col_keys)
            items.append(item_data)

    for bibliography_file in files:
        if bibliography_file.suffix.lower() not in BIBLIOGRAPHY_EXTENSIONS:
            continue
        col_keys = collection_keys_for(bibliography_file)
        for item_data in reference_entries_by_file.get(bibliography_file, []):
            if item_data["item_key"] in matched_reference_keys:
                continue
            item_data["collection_keys"] = json.dumps(col_keys)
            items.append(item_data)

    return items, collections
