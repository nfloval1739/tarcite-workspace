"""Document export routes."""

import io
from datetime import datetime
from typing import Any, Dict, List

from fastapi import APIRouter
from fastapi.responses import StreamingResponse

router = APIRouter(tags=["export"])


@router.post("/api/export/annotations/docx")
def export_annotations_docx(body: Dict[str, Any]):
    """Generate a .docx from the client-filtered annotations and theme tree."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    annotations = body.get("annotations") or []
    theme_tree = body.get("themes") or []

    doc = Document()

    title_para = doc.add_heading("Annotations Export", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Exported {datetime.now().strftime('%Y-%m-%d %H:%M')}  \u00b7  {len(annotations)} annotation(s)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    def _hex_to_rgb(hex_color: str) -> RGBColor:
        hex_value = (hex_color or "#3b82f6").lstrip("#")
        if len(hex_value) == 3:
            hex_value = "".join(c * 2 for c in hex_value)
        if len(hex_value) != 6:
            return RGBColor(0x3B, 0x82, 0xF6)
        return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))

    def _write_annotations(ann_list: List[Dict]):
        for annotation in ann_list:
            quote = (annotation.get("quote") or "").strip()
            comment = (annotation.get("comment") or "").strip()
            page = annotation.get("page") or ((annotation.get("page_index") or 0) + 1)
            source = annotation.get("source") or annotation.get("item_title") or annotation.get("item_key") or ""
            year = annotation.get("year") or annotation.get("item_year") or ""

            if quote:
                paragraph = doc.add_paragraph(style="Quote")
                run = paragraph.add_run(f'"{quote}"')
                run.italic = True

            if comment:
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"Note: {comment}")

            ref_parts = [source, str(year) if year else "", f"p. {page}"]
            ref_text = " \u00b7 ".join(x for x in ref_parts if x)
            ref_para = doc.add_paragraph()
            ref_run = ref_para.add_run(ref_text)
            ref_run.font.size = Pt(9)
            ref_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            doc.add_paragraph()

    def _write_theme_section(node: Dict, level: int):
        heading_level = min(level + 1, 4)
        color_hex = node.get("color") or "#3b82f6"
        heading = doc.add_heading(f"# {node['name']}", level=heading_level)
        for run in heading.runs:
            try:
                run.font.color.rgb = _hex_to_rgb(color_hex)
            except Exception:
                pass

        node_name = node["name"].lower()
        tag_anns = [
            annotation for annotation in annotations
            if any(t.get("name", "").lower() == node_name for t in (annotation.get("themes") or []))
        ]
        if tag_anns:
            _write_annotations(tag_anns)
        else:
            doc.add_paragraph()

        for child in node.get("children") or []:
            _write_theme_section(child, level + 1)

    if theme_tree:
        doc.add_heading("Themes", level=1)
        for root in theme_tree:
            _write_theme_section(root, level=2)

    untagged = [annotation for annotation in annotations if not annotation.get("themes")]
    if untagged:
        doc.add_heading("Untagged Annotations", level=1)
        _write_annotations(untagged)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"annotations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/api/export/themes/docx")
def export_themes_docx(body: Dict[str, Any]):
    """Generate a theme-organised .docx report from the client-sent tree + annotations."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    theme_tree: List[Dict] = body.get("themes") or []

    doc = Document()

    title_para = doc.add_heading("Themes Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Exported {datetime.now().strftime('%Y-%m-%d %H:%M')}  \u00b7  {len(theme_tree)} top-level theme(s)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    def _hex_to_rgb(hex_color: str) -> RGBColor:
        hex_value = (hex_color or "#3b82f6").lstrip("#")
        if len(hex_value) == 3:
            hex_value = "".join(c * 2 for c in hex_value)
        if len(hex_value) != 6:
            return RGBColor(0x3B, 0x82, 0xF6)
        return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))

    def _write_theme(node: Dict, level: int):
        heading_level = min(level, 4)
        heading = doc.add_heading(f"# {node['name']}", level=heading_level)
        for run in heading.runs:
            try:
                run.font.color.rgb = _hex_to_rgb(node.get("color") or "#3b82f6")
            except Exception:
                pass

        files = node.get("files", 0)
        refs = node.get("references", 0)
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"{files} file(s)  \u00b7  {refs} reference(s)")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        for annotation in node.get("annotations") or []:
            quote = (annotation.get("quote") or "").strip()
            comment = (annotation.get("comment") or "").strip()
            page = annotation.get("page", 1)
            source = annotation.get("source") or ""
            year = annotation.get("year") or ""

            if quote:
                paragraph = doc.add_paragraph(style="Quote")
                run = paragraph.add_run(f'"{quote}"')
                run.italic = True
            if comment:
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"Note: {comment}")

            ref_parts = [source, str(year) if year else "", f"p. {page}"]
            ref_text = " \u00b7 ".join(x for x in ref_parts if x)
            ref_para = doc.add_paragraph()
            ref_run = ref_para.add_run(ref_text)
            ref_run.font.size = Pt(9)
            ref_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            doc.add_paragraph()

        if not (node.get("annotations")):
            doc.add_paragraph()

        for child in node.get("children") or []:
            _write_theme(child, level + 1)

    for root in theme_tree:
        _write_theme(root, level=1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"themes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
