"""File content serving routes."""

import csv
import mimetypes
from pathlib import Path
from typing import Dict
from urllib.parse import quote

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse

from app.database import get_item_v2

router = APIRouter(tags=["content"])


@router.get("/api/pdf/{item_key}")
def serve_pdf(item_key: str):
    item = get_item_v2(item_key)
    if not item:
        raise HTTPException(status_code=404, detail="Item not found.")

    file_path_str = None
    files = item.get("files") or []
    if files:
        file_path_str = files[0].get("file_path")
    if not file_path_str:
        file_path_str = item.get("file_path")
    if not file_path_str:
        raise HTTPException(status_code=404, detail="No file associated with this item.")

    file_path = Path(file_path_str)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"File not found on disk: {file_path.name}")
    media_type = mimetypes.guess_type(file_path.name)[0] or "application/octet-stream"
    encoded_name = quote(file_path.name, safe="")
    return FileResponse(
        str(file_path),
        media_type=media_type,
        headers={"Content-Disposition": f"inline; filename*=UTF-8''{encoded_name}"},
    )


@router.get("/api/file-content/{item_key}")
def file_content_route(item_key: str) -> Dict:
    """Return structured text content for txt/md/csv/docx files."""
    item = get_item_v2(item_key)
    if not item:
        raise HTTPException(status_code=404, detail="Item not found.")

    files = item.get("files") or []
    file_path_str = files[0].get("file_path") if files else None
    if not file_path_str:
        file_path_str = item.get("file_path")
    if not file_path_str:
        raise HTTPException(status_code=404, detail="No file.")

    file_path = Path(file_path_str)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk.")

    ext = file_path.suffix.lower()

    if ext == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return {"type": "text", "content": text, "title": item.get("title", file_path.stem)}

    if ext in (".md", ".markdown"):
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return {"type": "markdown", "content": text, "title": item.get("title", file_path.stem)}

    if ext == ".csv":
        text = file_path.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(text.splitlines())
        rows = list(reader)
        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []
        return {
            "type": "csv",
            "title": item.get("title", file_path.stem),
            "headers": headers,
            "rows": data_rows,
            "total_rows": len(data_rows),
        }

    if ext == ".docx":
        try:
            from docx import Document

            doc = Document(str(file_path))
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    style = paragraph.style.name if paragraph.style else "Normal"
                    paragraphs.append({"text": paragraph.text, "style": style})
            tables = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    table_rows.append([cell.text.strip() for cell in row.cells])
                tables.append(table_rows)
            return {
                "type": "docx",
                "title": item.get("title", file_path.stem),
                "paragraphs": paragraphs,
                "tables": tables,
            }
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Cannot read docx: {exc}")

    raise HTTPException(status_code=415, detail=f"Unsupported type: {ext}")
